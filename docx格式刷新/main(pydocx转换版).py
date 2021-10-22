import os.path
from win32com import client
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pypandoc

NEW_FILE_LIST = []  # 临时文件列表
BASE_PATH = os.getcwd()  # 当前文件目录


def doc_docx(old_path):
    """单个.doc文件转换为.docx文件"""
    # 我们首先调用win32com模块并打开word应用软件，执行程序：

    new_path = old_path.split('.', 1)[0] + '.docx'
    print('创建临时文件',new_path)
    pypandoc.convert_file(old_path, 'doc', outputfile=new_path)

    return new_path


def int_docx():
    """对要输出的docx文件初始化"""
    new_docx = Document()
    # 纸张大小
    new_docx.sections[0].page_height = Cm(29.7)
    new_docx.sections[0].page_width = Cm(21)
    # 页边距
    new_docx.sections[0].top_margin = Cm(3.7)
    new_docx.sections[0].bottom_margin = Cm(3.4)
    new_docx.sections[0].left_margin = Cm(2.8)
    new_docx.sections[0].right_margin = Cm(2.6)
    # 字体大小三号字（16）
    new_docx.styles['Normal'].font.size = Pt(16)
    # 字体仿宋_GB2312
    new_docx.styles['Normal'].font.name = u'仿宋_GB2312'
    new_docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    # 行间距 28磅 段前、段后不留空
    new_docx.styles['Normal'].paragraph_format.line_spacing = Pt(29)
    new_docx.styles['Normal'].paragraph_format.space_before = Pt(0)
    new_docx.styles['Normal'].paragraph_format.space_after = Pt(0)
    # 首行缩进2字符
    new_docx.styles['Normal'].paragraph_format.first_line_indent = 406400
    # 关闭孤行控制
    # self.docx.styles['Normal'].paragraph_format.keep_together = True
    new_docx.styles['Normal'].paragraph_format.widow_control = False
    new_docx.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return new_docx


def docx_format_01(file_path):
    """对传入的word文档格式进行调整，并保存为新文档"""
    d = Document(file_path)
    new_docx = int_docx()  # 获取初始化过的docx对象

    # 字体大小

    for i in d.paragraphs:
        if len(i.text) == 0 or i.text.isspace():
            continue
        if i.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:  # 居中的按标题算
            print(i.text,'居中')
            # 首行处理（标题行）
            # 居中 不缩进
            p = new_docx.add_paragraph('')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.first_line_indent = 0
            # 字体、字号调整
            for r in i.runs:
                new_r = p.add_run(r.text)
                new_r.font.name = u'方正小标宋简体'
                new_r._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
                new_r.font.bold = False
                new_r.font.size = Pt(22)
            # 插入一个空行
            new_docx.add_paragraph('')
        else:
            new_p = new_docx.add_paragraph('')
            # 块 字体 字号
            for r in i.runs:
                new_r = new_p.add_run(r.text)
                new_r.font.bold = r.font.bold



    # 保存
    now_file_dir = os.path.join(BASE_PATH, '转换后')
    now_file = os.path.join(now_file_dir, os.path.basename(file_path))
    print(os.path.basename(file_path), '转换成功')
    if file_path in NEW_FILE_LIST:
        os.remove(file_path)
    new_docx.save(now_file)


def path_handle(path):
    """分析当前文件夹中的word文档，返回文档路径列表"""
    file_list = []
    for file in os.listdir(path):
        if file.split('.')[-1] == 'docx':
            file_list.append(os.path.join(path, file))
        elif file.split('.')[-1] == 'doc':
            new_path = doc_docx(os.path.join(path, file))
            file_list.append(new_path)
            NEW_FILE_LIST.append(new_path)

    return file_list


def run():
    print('当前路径', BASE_PATH)
    file_list = path_handle(BASE_PATH)
    print(f'共有{len(file_list)}个文件需要转换')
    if len(file_list) != 0:
        if not os.path.exists(os.path.join(BASE_PATH, '转换后')):
            os.mkdir(os.path.join(BASE_PATH, '转换后'))
        for i in file_list:
            docx_format_01(i)

    input_info = input('任意键退出\n')
    if input_info:
        return


if __name__ == '__main__':
    run()
