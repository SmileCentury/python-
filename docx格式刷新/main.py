import os.path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION

BASE_PATH = os.path.dirname(os.path.abspath(__file__))
def docx_format_01(file_path):
    """对传入的word文档格式进行调整，并保存为新文档"""
    d = Document(file_path)
    # 纸张大小
    d.sections[0].page_height = Cm(29.7)
    d.sections[0].page_width = Cm(21)
    # 页边距
    d.sections[0].top_margin = Cm(3.7)
    d.sections[0].bottom_margin = Cm(3.4)
    d.sections[0].left_margin = Cm(2.8)
    d.sections[0].right_margin = Cm(2.6)

    # 字体大小
    for i in d.paragraphs:
        # 关闭孤行控制
        i.paragraph_format.widow_control = False
        # 行 间距
        i.paragraph_format.line_spacing = Pt(29)
        i.paragraph_format.space_before = Pt(0)
        i.paragraph_format.space_after = Pt(0)
        # 块 字体 字号
        for r in i.runs:
            r.font.size = Pt(16)
            r.font.name = u'仿宋_GB2312'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

    # 首行处理
    #居中 不缩进
    d.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    d.paragraphs[0].paragraph_format.first_line_indent = 0
    # 字体、字号调整
    for r in d.paragraphs[0].runs:
        r.font.name = u'方正小标宋简体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
        r.font.bold = False
        r.font.size = Pt(22)
    # 插入一个空行
    d.paragraphs[1].insert_paragraph_before('')

    # 保存
    now_file_dir = os.path.join(BASE_PATH,'转换后')
    now_file = os.path.join(now_file_dir,os.path.basename(file_path))
    print(now_file)
    d.save(now_file)


def path_handle():
    """分析当前文件夹中的word文档，返回文档路径列表"""

    file_list = []
    for file in os.listdir(BASE_PATH):
        if file.split('.')[-1] in ['doc', 'docx']:
            file_list.append(os.path.join(BASE_PATH, file))

    return file_list

if __name__ == '__main__':
    file_list = path_handle()
    print(f'共有{len(file_list)}个文件需要转换')
    if len(file_list)!=0:
        if not os.path.exists(os.path.join(BASE_PATH,'转换后')):
            os.mkdir(os.path.join(BASE_PATH,'转换后'))
        for i in file_list:
            docx_format_01(i)